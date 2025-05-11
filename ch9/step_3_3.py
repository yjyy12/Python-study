from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm

from step_1_1 import IMG_DIR, OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_2_1 import OUT_2_1
from step_3_1 import apply_font
from step_3_2 import OUT_3_2, add_blank_paragraph

OUT_3_3 = OUT_DIR / f"{Path(__file__).stem}.docx"

def insert_indicators():
    doc = Document(OUT_3_2)  # 워드 문서 불러오기
    r_head = doc.add_paragraph().add_run("1. 주요 금리 현황")  # 부제목 입력
    apply_font(r_head, size_pt=14, is_bold=True)  # 폰트 설정
    add_blank_paragraph(doc, size_pt=10)  # 10 포인트 크기의 빈 단락 삽입

    table = doc.add_table(rows=1, cols=5)  # 표 추가(행 1개, 열 5개)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 표 가로 정렬
    table.allow_autofit = False  # 표 너비 자동 맞춤 해제

    tr = table.rows[0]  # 첫 번째 행 반환
    with pd.ExcelFile(OUT_2_1) as xlsx:  # ExcelFile 객체 생성
        for idx, sheet_name in enumerate(xlsx.sheet_names):  # 시트별로 반복 처리
            df_raw = pd.read_excel(xlsx, sheet_name=sheet_name)  # 데이터프레임 생성
            df_raw = df_raw.tail(24)  # 마지막 24개월 데이터 추출
            td = tr.cells[idx]  # tr 행에서 idx 위치의 셀 반환
            td.width = Mm(35.5)  # 셀 너비(35.5 밀리미터) 설정

            p1 = td.paragraphs[0]  # 첫 번째 단락 반환
            r1 = p1.add_run(sheet_name)  # 금리지표 이름 입력
            apply_font(r1, size_pt=12, is_bold=True, rgb="333333")

            p2 = td.add_paragraph()  # 단락 추가
            last_value = df_raw["DATA_VALUE"].iloc[-1]  # "DATA_VALUE" 열의 마지막 값
            r2 = p2.add_run(f"{last_value:,.2f}")  # 금리 입력
            apply_font(r2, size_pt=14, is_bold=True, rgb="333333")

            p3 = td.add_paragraph()  # 단락 추가
            diff = last_value - df_raw["DATA_VALUE"].iloc[0]  # 데이터 변동량
            arrow = "▲" if diff > 0 else "▼" if diff < 0 else ""  # 화살표
            rgb = "FF0000" if diff > 0 else "0000FF" if diff < 0 else "000000"  # 색상
            r3 = p3.add_run(f"{arrow}{abs(diff):,.2f}%p")  # 금리지표 변동량 입력
            apply_font(r3, size_pt=10, is_bold=True, rgb=rgb)

            p4 = td.add_paragraph()  # 단락 추가
            p4.paragraph_format.left_indent = Mm(-1)  # 단락 들여쓰기 설정
                        img_path = IMG_DIR / f"{sheet_name}.png"  # 이미지 경로
            p4.add_run().add_picture(img_path.as_posix(), Mm(30), Mm(8))  # 이미지 추가

    add_blank_paragraph(doc, size_pt=10)  # 10 포인트 크기의 빈 단락 삽입
    doc.save(OUT_3_3)  # 워드 파일로 저장

if __name__ == "__main__":
    insert_indicators()  # 주요 금리지표 데이터 입력