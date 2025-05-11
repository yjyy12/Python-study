from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_1_2 import OUT_1_2
from step_3_1 import apply_font
from step_3_2 import add_blank_paragraph
from step_3_3 import OUT_3_3

OUT_3_4 = OUT_DIR / f"{Path(__file__).stem}.docx"

def insert_deposit_info(n_rows: int = 10):
    doc = Document(OUT_3_3)  # 워드 문서 불러오기
    r_head = doc.add_paragraph().add_run("2. 주요 정기예금 상품 및 금리")  # 부제목 입력
    apply_font(r_head, size_pt=14, is_bold=True)  # 폰트 설정
    add_blank_paragraph(doc, size_pt=2)  # 2 포인트 크기의 빈 단락 삽입

    table = doc.add_table(rows=1, cols=6, style="Light Shading Accent 4")  # 표 추가
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 표 가로 정렬
    table.allow_autofit = False  # 표 너비 자동 맞춤 해제

    tr = table.rows[0]  # 첫 번째 행(헤더) 반환
    th_text = ["금융기관", "상품명", "이자계산", "만기(월)", "세전금리", "최고우대"]
    col_width = [Mm(40), Mm(53), Mm(20), Mm(20), Mm(20), Mm(20)]  # 열 너비
    for idx, th in enumerate(tr.cells):  # 열별로 반복 처리
        th.width = col_width[idx]  # 헤더 셀 너비 설정
        th.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 헤더 셀 세로 정렬
        p_th = th.paragraphs[0]  # 첫 번째 단락 반환
        p_th.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 헤더 셀 가로 정렬
        r_th = p_th.add_run(f"{th_text[idx]}")  # 헤더 셀 입력
        apply_font(r_th, size_pt=12, is_bold=True)

    df_raw = pd.read_excel(OUT_1_2)  # 정기예금 상품 데이터 불러오기
    df_filter = df_raw.filter(["kor_co_nm", "fin_prdt_nm", "intr_rate_type_nm", "save_trm", "intr_rate", "intr_rate2"])  # 사용할 열 필터링
    df_sort = df_filter.sort_values("intr_rate", ascending=False)  # 세전금리 순으로 정렬

    for _, se_row in df_sort.head(n_rows).iterrows():  # n_rows에 지정된 개수 반복 출력
        tr = table.add_row()  # 행 추가
        for idx, td in enumerate(tr.cells):  # 열별로 반복 처리
            td.width = col_width[idx]  # 내용 셀 너비 설정
            td.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 내용 셀 세로 정렬
            p_td = td.paragraphs[0]  # 첫 번째 단락 반환
            if idx < 2:  # 첫 두 개의 셀은 텍스트 배분 정렬 적용
                p_td.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            p_td.paragraph_format.space_before = Mm(2)  # 단락 앞 간격
            p_td.paragraph_format.space_after = Mm(2)  # 단락 뒤 간격
            p_td.add_run(f"{se_row.iloc[idx]}")  # 내용 셀 입력

    add_blank_paragraph(doc, size_pt=10)  # 10 포인트 크기의 빈 단락 삽입
    doc.save(OUT_3_4)  # 워드 파일로 저장

if __name__ == "__main__":
    insert_deposit_info(10)  # 정기예금 상품별 데이터 입력