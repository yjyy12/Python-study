from datetime import datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject  # Document 객체

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_3_1 import OUT_3_1, apply_font

OUT_3_2 = OUT_DIR / f"{Path(__file__).stem}.docx"

def add_blank_paragraph(doc: DocumentObject, size_pt: int = None):
    r_empty = doc.add_paragraph().add_run(" ")  # 빈 단락 삽입
    apply_font(r_empty, size_pt=size_pt)  # 폰트 설정

def add_title():
    doc = Document(OUT_3_1)  # 워드 문서 불러오기
    p_title = doc.add_paragraph(style="Title")  # 보고서 제목 삽입 및 "Title" 서식 적용
    r_title = p_title.add_run("정기예금 금리 현황표")  # 보고서 제목 입력
    apply_font(r_title, size_pt=20, is_bold=True)  # 폰트 설정

    now = datetime.now()  # 현재 시점의 datetime 객체 생성
    now_string = now.isoformat(sep=" ", timespec="minutes")  # 작성 일시 문자열 생성
    r_now = p_title.add_run(f" (작성 일시: {now_string})")  # 작성 일시 출력
    apply_font(r_now, size_pt=14)  # 폰트 설정
    add_blank_paragraph(doc, size_pt=5)  # 5 포인트 크기의 빈 단락 삽입
    doc.save(OUT_3_2)  # 워드 파일로 저장

if __name__ == "__main__":
    add_title()  # 보고서 제목 작성