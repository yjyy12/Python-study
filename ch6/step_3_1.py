from datetime import datetime as dt
from pathlib import Path

from docx import Document  # Document 함수 불러오기
from docx.document import Document as DocumentObject  # Document 객체
from docx.shared import Pt  # 폰트 크기 설정 클래스(포인트 단위)
from docx.text.run import Run

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.


def apply_font_style(run: Run, size_pt: int = None, is_bold: bool = None):
    if size_pt is not None:
        run.font.size = Pt(size_pt)  # 폰트 크기 설정
    if is_bold is not None:
        run.font.bold = is_bold  # 폰트 두께 설정


def init_docx() -> DocumentObject:
    doc = Document()  # Document 객체 생성
    p1 = doc.add_paragraph(style="Heading 1")  # Paragraph 객체 생성
    run1 = p1.add_run("쇼핑 트렌드 보고서")  # Run 객체 생성
    apply_font_style(run1, size_pt=25, is_bold=True)
    date_string = dt.now().strftime(" (%Y.%m.%d)")  # 현재 날짜
    apply_font_style(p1.add_run(date_string), size_pt=15)
    return doc  # Document 객체 반환


if __name__ == "__main__":
    doc = init_docx()
    doc.save(OUT_DIR / f"{Path(__file__).stem}.docx")  # 워드 문서 파일로 저장