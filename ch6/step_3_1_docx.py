from docx import Document  # Document 함수 불러오기
from docx.shared import Pt  # Pt(포인트) 클래스 불러오기

doc = Document()  # Document 객체 생성
p1 = doc.add_paragraph()  # Paragraph 객체 생성
run1 = p1.add_run("Hello, ")  # Run 객체 생성
run1.font.size = Pt(25)  # 폰트 크기를 25 포인트로 설정

run2 = p1.add_run("World!")  # Run 객체 생성
run2.font.bold = True  # 폰트 두께를 두껍게 설정

doc.save("path_to_file.docx")  # 워드 문서 파일로 저장
