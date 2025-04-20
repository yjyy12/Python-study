from docx import Document  # Document 함수 불러오기

doc = Document()  # Document 객체 생성
table = doc.add_table(rows=3, cols=2, style="Table Grid")  # 표 생성
idx = 1
for tr in table.rows:  # 행 반복 처리
    for td in tr.cells:  # 열 반복 처리
        p_cell = td.paragraphs[0]  # 현재 셀의 기본 Paragraph 객체 선택
        p_cell.add_run(f"{idx}")  # Run 객체 생성
        idx += 1
doc.save("path_to_file.docx")
