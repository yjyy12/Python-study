import json
from pathlib import Path
from docx import Document  
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import asyncio

from step_1_1 import OUT_DIR  # 이전에 작성한 모듈을 불러옵니다.
from step_2_2 import OUT_2_2
from step_2_3 import fetch_trends_by_filter
from step_3_1 import apply_font_style, init_docx

async def add_table(doc: DocumentObject, category: str, option: str):
    await fetch_trends_by_filter(category, option)
    imgs_path: list = json.loads(OUT_2_2.read_text(encoding="utf8"))
    n_items = len(imgs_path)
    n_cols = 5
    n_rows = n_items // n_cols + (1 if n_items % n_cols > 0 else 0)

    para = doc.add_paragraph(style="Heading 2")
    text_filter = f"{option}의 {category} 트렌드"
    apply_font_style(para.add_run(text_filter), size_pt=15, is_bold=True)

    table = doc.add_table(rows=n_rows, cols=n_cols, style="Table Grid")
    for tr in table.rows:
        for td in tr.cells:
            if len(imgs_path) > 0:
                img_path = imgs_path.pop(0)
                p_cell = td.paragraphs[0]
                p_cell.add_run().add_picture(img_path, width=Cm(3))
                p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    p_list = doc.add_paragraph(style="List Bullet")
    text_notice = "보다 자세한 정보는 네이버플러스 스토어에서 확인하세요."
    apply_font_style(p_list.add_run(text_notice), size_pt=9)

async def main():
    doc = init_docx()  # 여기서 문서 생성
    await add_table(doc, "패션의류", "10대 여성")  # 표 추가
    doc.save(OUT_DIR / f"{Path(__file__).stem}.docx")  # 저장

if __name__ == "__main__":
    asyncio.run(main())

